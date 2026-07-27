import os
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from docx import Document
from PySide6.QtWidgets import QApplication

import preview_service
from preview_service import (
    docx_to_html,
    locate_docx_content_pages,
    prepare_preview,
    render_docx_with_office,
)


class PreviewServiceTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.application = QApplication.instance() or QApplication([])

    def test_offline_preview_and_page_location(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            directory = Path(temporary_directory)
            source = directory / "card.docx"
            document = Document()
            document.add_heading("工艺卡", level=1)
            document.add_paragraph("第一页工序内容：粗车外圆。")
            document.add_page_break()
            document.add_paragraph("第二页工序内容：钻孔并倒角。")
            document.save(source)

            with patch.dict(
                os.environ,
                {
                    "DOCSWIFT_PREVIEW_RENDERER": "html",
                    "LOCALAPPDATA": str(directory),
                },
            ):
                preview = prepare_preview(source)

                self.assertTrue(preview.pdf_path.exists())
                self.assertGreaterEqual(preview.page_count, 2)
                locations = locate_docx_content_pages(
                    source,
                    ["第一页工序内容：粗车外圆。", "第二页工序内容：钻孔并倒角。"],
                )
            self.assertEqual((1, 1), locations[0])
            self.assertEqual((2, 2), locations[1])

    def test_html_fallback_keeps_visible_table_borders(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            source = Path(temporary_directory) / "table.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "工序"
            table.cell(0, 1).text = "内容"
            table.cell(1, 0).text = "车"
            table.cell(1, 1).text = "粗车外圆"
            document.save(source)

            html = docx_to_html(source)
            self.assertIn('border="1"', html)
            self.assertIn("border:1px solid #333", html)

    def test_native_office_renderer_prefers_wps(self) -> None:
        with patch.object(
            preview_service,
            "_render_docx_with_com",
            return_value=True,
        ) as render:
            self.assertTrue(render_docx_with_office("source.docx", "preview.pdf"))
        render.assert_called_once_with(
            "source.docx",
            "preview.pdf",
            "kwps.Application",
        )

    def test_repeated_content_uses_operation_number_anchor(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            directory = Path(temporary_directory)
            source = directory / "repeated.docx"
            document = Document()
            document.add_paragraph("检验 11 检验")
            document.add_page_break()
            document.add_paragraph("检验 25 检验")
            document.save(source)

            with patch.dict(
                os.environ,
                {
                    "DOCSWIFT_PREVIEW_RENDERER": "html",
                    "LOCALAPPDATA": str(directory),
                },
            ):
                locations = locate_docx_content_pages(
                    source,
                    ["检验", "检验"],
                    ["11", "25"],
                    ["检验", "检验"],
                )
            self.assertEqual([(1, 1), (2, 2)], locations)


if __name__ == "__main__":
    unittest.main()
